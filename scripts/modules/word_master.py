# -*- coding: utf-8 -*-
"""
Genererer Word-masterplan-snapshot fra plan.json.

To dokumenter genereres:
  data/Fast_as_Fifty_Masterplan_2026.docx — Kennets 14 uger
  data/Eva_Medoc_Traningsplan_2026.docx    — Evas 13 uger

Struktur pr. plan:
  Title-side (program-navn, dato, race-mål)
  Én uge pr. side: uge-metadata (block, ctl-mål, TSS, location) + tabel over 7 dage
  Slut-side: full-plan summary

Kaldes fra apply_edit.py efter plan.json-commit. Onedrive-sync.yml
fanger commits automatisk.
"""
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

DAYS_DA = ["Man", "Tir", "Ons", "Tor", "Fre", "Lør", "Søn"]
MON_DA = ["jan", "feb", "mar", "apr", "maj", "jun", "jul", "aug", "sep", "okt", "nov", "dec"]

WINE = RGBColor(0x59, 0x18, 0x2A)
GOLD = RGBColor(0xBE, 0x86, 0x3A)
MUTED = RGBColor(0x7A, 0x6A, 0x58)
TEAL = RGBColor(0x1A, 0x7A, 0x5E)


def _iso(d: date) -> str:
    return d.isoformat()


def _add_days(d_iso: str, n: int) -> str:
    return _iso(date.fromisoformat(d_iso) + timedelta(days=n))


def _fmt_date(d_iso: str) -> str:
    d = date.fromisoformat(d_iso)
    return f"{d.day}. {MON_DA[d.month - 1]}"


def _fmt_dur(secs) -> str:
    if not secs:
        return ""
    m = int(secs) // 60
    if m >= 90 and m % 30 == 0:
        h, rem = divmod(m, 60)
        return f"{h}t{'' if rem == 0 else f' {rem}m'}"
    return f"{m} min"


def _run(paragraph, text, *, bold=False, size=11, color=None):
    r = paragraph.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def _title_page(doc, program, races, race_label, theme_color, is_eva=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(140)
    _run(p, "FAST AS FIFTY" if not is_eva else "EVA · MÉDOC",
         bold=True, size=32, color=theme_color)
    p = doc.add_paragraph()
    _run(p, program["name"], bold=True, size=18)
    p = doc.add_paragraph()
    _run(p, f"{program['totalWeeks']} uger · start {_fmt_date(program['start'])}",
         size=12, color=MUTED)
    p = doc.add_paragraph()
    _run(p, "Racemål", bold=True, size=13, color=theme_color)
    for r in races:
        p = doc.add_paragraph()
        _run(p, f"{_fmt_date(r['date'])}  ·  {r['name']}", size=12)
    p = doc.add_paragraph()
    _run(p, race_label, size=10, color=MUTED)
    doc.add_page_break()


def _week_page(doc, week_meta: dict, days_data: list, note_key: str,
               theme_color, entry_zones=None):
    """Én side pr. uge."""
    p = doc.add_paragraph()
    _run(p, f"Uge {week_meta['week']}", bold=True, size=22, color=theme_color)

    p = doc.add_paragraph()
    sub_parts = []
    if week_meta.get("blockType"):
        sub_parts.append(week_meta["blockType"])
    if week_meta.get("location"):
        sub_parts.append(week_meta["location"])
    if week_meta.get("ctlTarget"):
        sub_parts.append(f"CTL-mål {week_meta['ctlTarget']}")
    if week_meta.get("tssTarget"):
        sub_parts.append(f"{week_meta['tssTarget']} TSS")
    _run(p, "  ·  ".join(sub_parts), size=11, color=MUTED)

    if week_meta.get(note_key):
        p = doc.add_paragraph()
        _run(p, week_meta[note_key], size=11)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    hdr[0].text = "Dag"; hdr[1].text = "Pas"; hdr[2].text = "Detaljer"
    for c in hdr:
        for r in c.paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(11)

    for day_iso in [_add_days(week_meta["start"], i) for i in range(7)]:
        row = tbl.add_row().cells
        d = date.fromisoformat(day_iso)
        row[0].text = f"{DAYS_DA[(d.weekday())]} {_fmt_date(day_iso)}"
        day = next((x for x in days_data if x["date"] == day_iso), None)
        if not day or not day.get("entries") or all(not e.get("workout") for e in day["entries"]):
            note = ""
            if day and day.get("entries"):
                note = " · ".join(e.get("note", "") for e in day["entries"] if e.get("note")) or "Hviledag"
            else:
                note = "Hviledag"
            row[1].text = "Hvile"
            row[2].text = note
        else:
            names = []
            details = []
            for e in day["entries"]:
                wo = e.get("workout")
                if not wo:
                    if e.get("note"):
                        details.append(e["note"])
                    continue
                nm = wo.get("name", "")
                if wo.get("moving_time"):
                    nm += f" · {_fmt_dur(wo['moving_time'])}"
                names.append(nm)
                if wo.get("description"):
                    details.append(wo["description"])
                if e.get("note"):
                    details.append(e["note"])
            row[1].text = "\n".join(names)
            row[2].text = "\n\n".join(details)
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    tbl.columns[0].width = Cm(3)
    tbl.columns[1].width = Cm(6)
    tbl.columns[2].width = Cm(9)
    doc.add_page_break()


def _summary_page(doc, program, weeks_meta, note_key, theme_color):
    p = doc.add_paragraph()
    _run(p, "Overblik", bold=True, size=18, color=theme_color)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Uge"; hdr[1].text = "Fase"; hdr[2].text = "Note"
    for c in hdr:
        for r in c.paragraphs[0].runs:
            r.font.bold = True
    for w in weeks_meta:
        row = tbl.add_row().cells
        row[0].text = f"Uge {w['week']} — {_fmt_date(w['start'])}"
        row[1].text = w.get("blockType", "")
        row[2].text = w.get(note_key, "") or w.get("location", "")


def _season_page(doc, season, theme_color):
    """Faseoversigt for naeste saeson. Ikke ugesider — sessioner planlaegges
    foerst naar hver fase aabner."""
    doc.add_page_break()
    p = doc.add_paragraph()
    _run(p, f"SÆSON {season['year']}", bold=True, size=18, color=theme_color)

    p = doc.add_paragraph()
    _run(p, f"A-løb: Tour des Stations Ultrafondo · 28. aug 2027 · 242 km / 8.848 hm",
         bold=True, size=12)
    p = doc.add_paragraph()
    _run(p, f"CTL {season['ctlAtRace']} på startlinjen (peak {season['ctlPeak']}) · "
            f"FTP {season['ftpStart']} → {season['ftpTarget']} · "
            f"{season['hoursPerWeekAvg']} t/uge i snit, {season['hoursPerWeekPeak']} i toppen",
         size=10, color=MUTED)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Fase", "Uger", "Periode", "CTL", "FTP"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True

    weeks = {w["week"]: w for w in season["weeks"]}
    for ph in season["phases"]:
        a, b = ph["weekFrom"], ph["weekTo"]
        wa, wb = weeks[a], weeks[b]
        d1 = date.fromisoformat(wa["start"])
        d2 = date.fromisoformat(wb["start"]) + timedelta(days=6)
        row = tbl.add_row().cells
        row[0].text = ph["name"]
        row[1].text = f"{a}-{b}"
        row[2].text = (f"{d1.day}. {MON_DA[d1.month-1]} {str(d1.year)[2:]} → "
                       f"{d2.day}. {MON_DA[d2.month-1]} {str(d2.year)[2:]}")
        row[3].text = f"{wa['ctlTarget']}→{wb['ctlTarget']}"
        row[4].text = str(ph["ftpTarget"])

    for key, label in [("CAMP", "Lejre — her bygges CTL"), ("RACE", "Løb")]:
        hits = [w for w in season["weeks"] if w["blockType"] == key]
        if not hits:
            continue
        p = doc.add_paragraph()
        _run(p, label, bold=True, size=11, color=theme_color)
        for w in hits:
            p = doc.add_paragraph()
            _run(p, f"Uge {w['week']} · {_fmt_date(w['start'])} "
                    f"{date.fromisoformat(w['start']).year} · {w['location']} · "
                    f"CTL {w['ctlTarget']} · {w['tssTarget']} TSS", size=10)
            if w.get("note"):
                p = doc.add_paragraph()
                _run(p, f"    {w['note']}", size=9, color=MUTED)

    wp = season.get("weightPlan")
    if wp:
        p = doc.add_paragraph()
        _run(p, "Vægt", bold=True, size=11, color=theme_color)
        p = doc.add_paragraph()
        _run(p, f"{wp['startKg']} → {wp['targetKg']} kg, max {wp['maxLossPerWeekKg']} kg/uge. "
                f"Mål nået {wp['holdFromMonth']}, derefter vedligehold.", size=10)
        p = doc.add_paragraph()
        _run(p, f"    {wp['note']}", size=9, color=MUTED)


def generate_kennet(plan: dict) -> bytes:
    doc = Document()
    prog = plan["program"]
    races = plan["races"]
    kennet = plan["athletes"]["kennet"]
    stamp = f"Genereret {date.today().isoformat()} fra plan.json"

    _title_page(doc, prog, races, stamp, WINE, is_eva=False)
    for wm in plan["weeks"]:
        _week_page(doc, wm, kennet["days"], "note", WINE)
    _summary_page(doc, prog, plan["weeks"], "note", WINE)
    if plan.get("season2027"):
        _season_page(doc, plan["season2027"], GOLD)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_eva(plan: dict) -> bytes:
    doc = Document()
    eva = plan["athletes"]["eva"]
    prog = eva["program"]
    race = next((r for r in plan["races"] if r["date"] == prog["raceDay"]), None)
    races = [race] if race else []
    stamp = f"Genereret {date.today().isoformat()} fra plan.json"

    _title_page(doc, prog, races, stamp, TEAL, is_eva=True)
    for wm in eva["weeks"]:
        _week_page(doc, wm, eva["days"], "note", TEAL)
    _summary_page(doc, prog, eva["weeks"], "note", TEAL)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_both(plan: dict) -> dict[str, bytes]:
    return {
        "data/Fast_as_Fifty_Masterplan_2026.docx": generate_kennet(plan),
        "data/Eva_Medoc_Traningsplan_2026.docx": generate_eva(plan),
    }
